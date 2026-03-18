"""
Post-processors: convert rake analysis outputs to user-friendly formats.

The most common use-case is turning a Markdown report written by the agent
into a Word document (.docx) for easy sharing.

Usage::

    from rake_sdk.preprocessors.postprocessors import postprocess_files

    result = await client.analyze_bytes({"report.docx": doc_bytes})
    # result.output_files may contain "report.md" — convert back to .docx:
    deliverables = postprocess_files(result.output_files)
    # deliverables["report.docx"] is now ready for download
"""
from __future__ import annotations

import io
import re
from pathlib import Path


def postprocess_file(filename: str, content: bytes) -> dict[str, bytes]:
    """Post-process a single output file. Returns {filename: content, ...}."""
    ext = Path(filename).suffix.lower()
    if ext == ".md":
        docx_bytes = _md_to_docx(content)
        if docx_bytes is not None:
            docx_name = str(Path(filename).with_suffix(".docx"))
            return {filename: content, docx_name: docx_bytes}
    return {filename: content}


def postprocess_files(output_files: dict[str, bytes]) -> dict[str, bytes]:
    """Post-process all output files from a rake analysis run."""
    result: dict[str, bytes] = {}
    for filename, content in output_files.items():
        result.update(postprocess_file(filename, content))
    return result


# ── md → docx ────────────────────────────────────────────────────────────────

def _md_to_docx(content: bytes) -> bytes | None:
    """Convert Markdown bytes to a DOCX document. Returns None if python-docx is not installed."""
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError:
        return None

    text = content.decode("utf-8", errors="replace")
    doc = Document()

    # Remove default empty paragraph Word adds
    for para in list(doc.paragraphs):
        p = para._element
        p.getparent().remove(p)

    _heading_re = re.compile(r"^(#{1,6})\s+(.+)")
    _bullet_re  = re.compile(r"^[-*]\s+(.+)")
    _num_re     = re.compile(r"^\d+\.\s+(.+)")
    _hr_re      = re.compile(r"^---+$|^\*\*\*+$")
    _code_fence = re.compile(r"^```")

    in_code_block = False
    code_lines: list[str] = []

    def flush_code() -> None:
        if code_lines:
            p = doc.add_paragraph("\n".join(code_lines), style="No Spacing")
            p.runs[0].font.name = "Courier New"
            p.runs[0].font.size = Pt(9)
        code_lines.clear()

    for line in text.splitlines():
        if _code_fence.match(line):
            if in_code_block:
                flush_code()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        m_heading = _heading_re.match(line)
        m_bullet  = _bullet_re.match(line)
        m_num     = _num_re.match(line)

        if m_heading:
            level = min(len(m_heading.group(1)), 9)
            doc.add_heading(_strip_inline_md(m_heading.group(2)), level=level)
        elif m_bullet:
            doc.add_paragraph(_strip_inline_md(m_bullet.group(1)), style="List Bullet")
        elif m_num:
            doc.add_paragraph(_strip_inline_md(m_num.group(1)), style="List Number")
        elif _hr_re.match(line.strip()):
            # Horizontal rule → paragraph border would need XML; just add spacing
            doc.add_paragraph("")
        elif line.strip():
            doc.add_paragraph(_strip_inline_md(line))
        else:
            doc.add_paragraph("")

    if in_code_block:
        flush_code()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_INLINE_RE = re.compile(
    r"\*\*(.+?)\*\*"    # bold
    r"|\*(.+?)\*"        # italic
    r"|`(.+?)`"          # inline code
    r"|\[(.+?)\]\(.+?\)" # link
)


def _strip_inline_md(text: str) -> str:
    """Strip simple inline markdown markers, keeping the visible text."""
    def repl(m: re.Match) -> str:
        return m.group(1) or m.group(2) or m.group(3) or m.group(4) or ""
    return _INLINE_RE.sub(repl, text)
